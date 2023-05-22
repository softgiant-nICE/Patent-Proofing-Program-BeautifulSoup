from urllib.request import urlopen
from bs4 import BeautifulSoup
import requests
import os
import re
import glob
import pandas as pd
import win32com.client
import docx
from docx import Document
import pandas as pd
import numpy as np
import os
from datetime import datetime
import re
from docx import Document
import re
from datetime import datetime, timedelta
from dateutil.relativedelta import relativedelta
from datetime import date
from docx.text.run import *
import shutil
################################################################################
################################################################################
                            # Functions
################################################################################
# Function to delete files in folder, used to clean the directory of previous docs
def clean_dir(x):
    try:
        os.remove(x)
    except:
        pass
    return

# Function to create a new word document from a string, and save as filenamepython
def new_docx(filename, *texts):
    document = Document()
    for text in texts:
        p = document.add_paragraph(text)
        run = p.add_run()
        run.add_break(docx.enum.text.WD_BREAK.PAGE)
    document.save(filename)

# Function to create a spreadsheet with the cited and published references
def new_spreadsheet(filename, ipo_us, ipo_for, ipo_npl, bip_us, bip_for, bip_npl):
    import openpyxl
    from openpyxl import Workbook
    from openpyxl.utils import FORMULAE
    wb = Workbook()
    ws1 = wb.create_sheet('Patent References')
    ws1['B1'] = 'Cited'
    ws1['C1'] = 'Printed'
    ws1['D1'] = 'Cited Not Printed'
    ws1['E1'] = 'Printed Not Cited'
    ws2 = wb.create_sheet('NPL References')
    ws2['B1'] = 'Cited'
    ws2['D1'] = 'Printed'
    i = 2
    e = 2
    ### List US patent references on Worksheet Patent References ###
    # What was printed on the patent
    for n in ipo_us.split('\n'):
        cell = 'C' + str(i)
        ws1[cell] = n
        i += 1
    # What was cited by BIP
    for line in bip_us:
        cell = 'B' + str(e)
        e += 1
        ws1[cell] = line
    ### Now list Foreign Patent References on Worksheet Patent References ###
    # What was printed on the patent
    for n in ipo_for.split('\n'):
        cell = 'C' + str(i)
        ws1[cell] = n
        i += 1
    # What was cited by BIP    
    for line in bip_for:
        cell = 'B' + str(e)
        e += 1
        ws1[cell] = line
    ### What was cited but not printed and what was printed but not cited ###
    # Cited but not printed NOTE: FILTER function not available in openpyxlquit. This is why = is not included in formula
    ws1['D2'] = 'FILTER(B2:B' + str(e) + ',ISNA(VLOOKUP(TRIM(B2:B' + str(e) + '),TRIM(C2:C' + str(i) + '),1,FALSE)))'
    # Printed but not cited
    ws1['E2'] = 'FILTER(C2:C' + str(i) + ',ISNA(VLOOKUP(TRIM(C2:C' + str(i) + '),TRIM(B2:B' + str(e) + '),1,FALSE)))'
    ### We are going to put NPLs on second Worksheet NPL References ###
    i = 2
    e = 2
    # What was printed on the patent
    for line in ipo_npl:
        cell = 'D' + str(i)
        ws2[cell] = line
        i += 1
    # What was cited by BIP
    for line in bip_npl:
        cell = 'B' + str(e)
        e += 1
        ws2[cell] = line
    wb.save(filename)

# Function to compare docs
def compare(file_1, file_2, output):
    #Create the Application word, https://stackoverflow.com/questions/47212459/automating-comparison-of-word-documents-using-python
    path = os.getcwd()
    Application=win32com.client.gencache.EnsureDispatch("Word.Application")
    # Compare documents
    Application.CompareDocuments(Application.Documents.Open(os.getcwd() + file_1), Application.Documents.Open(os.getcwd() + file_2))
    # Save the comparison document
    Application.ActiveDocument.ActiveWindow.View.Type = 3
    Application.ActiveDocument.SaveAs(FileName = path + "\\" + output)
    Application.Quit()

# Function to extract text and remove formatting from word doc.
def get_text(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

# Clean the directory of old comparison files and downloads to ensure you start fresh
# def del_prev_docs():
    # clean_dir('PTO Claims.docx')
    # clean_dir('PTO Description.docx')
    # clean_dir('PTO References.docx')
    # clean_dir('Compare Description.docx')
    # clean_dir('Compare Claims.docx')
    # clean_dir('Compare References.docx')
    # clean_dir('BIP Claims.docx')
    # clean_dir('BIP Description.docx')
    # clean_dir('BIP References.docx')

################################################################################
################################################################################
            # Download PTO's issued patent version of documents from Google
################################################################################


def pto_docs(pat_no):
    global description
    # Use the directory to specify the patent number, then get the patent text
    # from Google Patents.  Header info is from https://python-forum.io/thread-28613.html
    url = 'https://patents.google.com/patent/US{0}/en?oq= {0}'.format(pat_no)
    #headers = {
    #    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/84.0.4147.89 Safari/537.36'
    #}
    global soup
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'lxml')
    # Pandas has a way to pull data from tables on a website.  Using pandas to
    # grab the US and forn patent references and NPLs quickly.  The 8th (0-indexed, so
    # 7) table is the US dofor_cs, 9th is foreign refs, 10th is NPLs.  Includes some formatting
    pd.set_option("display.max_colwidth", 10000)
    df_list = pd.read_html(response.text)
    us_pat_docs = ''
    for_pat_docs = ''
    npl_docs = ''
    count = 1
    for df in df_list:
        ###### Patent Docs ########
        ###### This is to account for the tables not being the same ######
        ###### for every patent                                     ######
        if 'Country' in str(df.columns):
            count = 1
        elif 'Assignee' in str(df.columns):
            count += 1
            if count == 3:
                us_pat_docs = df[df['Publication number'].str.contains('US')]
                for_pat_docs = df[df['Publication number'].str.contains('US') == False]
        ####### NPL Docs ########
        elif 'Title' in str(df.columns) and 'Application Number' not in str(df.columns) and 'Publication number' not in str(df.columns) and 'Publication' not in str(df.columns) and 'Code' not in str(df.columns):
            npl_docs = df
    ###### US Patent Docs ########
    if str(us_pat_docs) != '':
        us_pat_docs = us_pat_docs[us_pat_docs.iloc[:,1].map(lambda x: x==x)]
        us_pat_docs = us_pat_docs.reset_index(drop=True)
        us_pat_docs.drop(us_pat_docs.head(1).index,inplace=True)
        us_pat_docs['Publication number'] = us_pat_docs['Publication number'].map(lambda x: re.sub('A1','',str(x)))
        us_pat_docs['Publication number'] = us_pat_docs['Publication number'].map(lambda x: re.sub('A2','',str(x)))
        us_pat_docs['Publication number'] = us_pat_docs['Publication number'].map(lambda x: re.sub('B1','',str(x)))
        us_pat_docs['Publication number'] = us_pat_docs['Publication number'].map(lambda x: re.sub('B2','',str(x)))
        us_pat_docs['Publication number'] = us_pat_docs['Publication number'].map(lambda x: re.sub('[a-zA-Z()*]','',str(x)))
        us_pat_docs['Publication number'] = us_pat_docs['Publication number'].map(lambda x: int(x))
        us_pat_docs = us_pat_docs.sort_values(by=['Publication number']).reset_index(drop=True)
        us_pat_docs = us_pat_docs['Publication number']
        us_pat_docs = us_pat_docs.to_string(index=False,header=False)
    if str(for_pat_docs) != '':
        for_pat_docs = for_pat_docs[for_pat_docs.iloc[:,1].map(lambda x: x==x)]
        for_pat_docs['Publication number'] = for_pat_docs['Publication number'].map(lambda x: re.sub('A1','',str(x)))
        for_pat_docs['Publication number'] = for_pat_docs['Publication number'].map(lambda x: re.sub('A2','',str(x)))
        for_pat_docs['Publication number'] = for_pat_docs['Publication number'].map(lambda x: re.sub('B1','',str(x)))
        for_pat_docs['Publication number'] = for_pat_docs['Publication number'].map(lambda x: re.sub('A','',str(x)))
        for_pat_docs['Publication number'] = for_pat_docs['Publication number'].map(lambda x: re.sub('B','',str(x)))
        for_pat_docs['Publication number'] = for_pat_docs['Publication number'].map(lambda x: str(x[ : -6]))
        for_pat_docs = for_pat_docs.sort_values(by=['Publication number']).reset_index(drop=True)
        for_pat_docs = for_pat_docs['Publication number']
        for_pat_docs = for_pat_docs.to_string(index=False,header=False)
    ####### NPL Docs ##########
    if str(npl_docs) != '':
        npl_docs = npl_docs[npl_docs.iloc[:,0].map(lambda x: x==x)]
        npl_docs = npl_docs['Title']
        npl_docs = npl_docs.to_string(index=False,header=False)
    #references = us_pat_docs + '\n' + for_pat_docs + '\n\n\n\n\n\n\n\n\n' + npl_docs
    # The patent sections are split out by flags.  This finds those flags and isolates
    # the text we want from the patent.
    a = str(soup).split('<h2>Abstract</h2>')[1]
    abstract = a.split('</abstract>')[0]
    b = a.split('<h2>Description</h2>')[1]
    priority = b.split('BACKGROUND</heading>')[0]
    c = b.split('BACKGROUND</heading>')[1]
    description = c.split('<section itemprop="claims" itemscope="">')[0]
    d = c.split('<section itemprop="claims" itemscope="">')[1] 
    claims = d.split('</section>')[0]
    e = str(soup).split('<head>')[1]
    title = e.split('<meta content="width=device-width, initial-scale=1" name="viewport"/>')[0]

    title = BeautifulSoup(title, features = 'lxml').get_text()
    priority = BeautifulSoup(priority, features = 'lxml').get_text()
    abstract = BeautifulSoup(abstract, features = 'lxml').get_text()
    description = BeautifulSoup(description, features = 'lxml').get_text()
    claims = BeautifulSoup(claims, features = 'lxml').get_text()

    return claims, title+priority+'\n\n\nBACKGROUND\n'+description+'\n\n\nAbstract\n'+abstract, us_pat_docs, for_pat_docs, npl_docs

def new_func(df):
    return str(df.columns)


################################################################################
################################################################################
                # Get BIP's as-filed versions of docs
################################################################################
# Expecting the IDS form to be saved in a subfolder as 'IDS - Forms.csv'
def bip_docs():
    global description
    subfolders = [ f.path for f in os.scandir(os.getcwd()) if f.is_dir() ]
    pd.set_option("display.max_colwidth", 10000)
    for subfolder in subfolders:
        if 'IDS' in subfolder:
            try:
                refs = pd.read_csv(subfolder + '\\IDS - Forms.csv', encoding = 'unicode_escape')
            except:
                pass
            try:
                refs = pd.read_csv(subfolder + '\\report.csv', encoding = 'unicode_escape')
            except:
                pass
    # Breaking out the foreign, npl, US issued patents, and US patent pubs
    forn_doc_cols = [x for x in refs.columns if 'foreign' in x and 'relevant' not in x]
    npl_doc_cols = [x for x in refs.columns if 'nplcit' in x]
    uspat_doc_cols = [x for x in refs.columns if 'us-patent-cite' in x]
    uspub_doc_cols = [x for x in refs.columns if 'us-pub-appl' in x]
    # US Patent Docs
    uspat_docs = refs[uspat_doc_cols]
    uspat_refs = pd.DataFrame([['','','']],columns=['Ref No', 'Date', 'Name'])
    i = 0
    while True:
        doc_no = 'us-ids[0].us-patent-cite[0].us-doc-reference[{}].doc-number[0]'.format(i)
        date = 'us-ids[0].us-patent-cite[0].us-doc-reference[{}].date[0]'.format(i)
        name = 'us-ids[0].us-patent-cite[0].us-doc-reference[{}].name[0]'.format(i)
        try:
            ids = uspat_docs[[doc_no, date, name]]
            ids = ids.rename(columns={doc_no:'Ref No', date:'Date', name:'Name'})
            ids = ids[ids.iloc[:,1].map(lambda x: x==x)]
            uspat_refs = pd.concat([uspat_refs,ids])
            i += 1
        except KeyError:
            break
            #
            #
            #
    uspat_refs = uspat_refs.reset_index(drop=True)
    uspat_refs.drop(uspat_refs.head(1).index,inplace=True)
    uspat_refs['Ref No'] = uspat_refs['Ref No'].map(lambda x: int(x))
    uspat_refs = uspat_refs.sort_values(by=['Ref No']).reset_index(drop=True)
    try:
        uspat_refs['Date'] = uspat_refs['Date'].map(lambda x: datetime.strptime(str(x),'%Y-%m-%d').strftime('%B %Y'))
    except ValueError:
        uspat_refs['Date'] = uspat_refs['Date'].map(lambda x: datetime.strptime(str(x),'%m/%d/%Y').strftime('%B %Y'))
    uspub_docs = refs[uspub_doc_cols]
    uspub_refs = pd.DataFrame([['','','']],columns=['Ref No', 'Date', 'Name'])
    i = 0
    while True:
        doc_no =  'us-ids[0].us-pub-appl-cite[0].us-doc-reference[{}].doc-number[0]'.format(i)
        date = 'us-ids[0].us-pub-appl-cite[0].us-doc-reference[{}].date[0]'.format(i)
        name = 'us-ids[0].us-pub-appl-cite[0].us-doc-reference[{}].name[0]'.format(i)
        try:
            ids = uspub_docs[[doc_no, date, name]]
            ids = ids.rename(columns={doc_no:'Ref No', date:'Date', name:'Name'})
            ids = ids[ids.iloc[:,1].map(lambda x: x==x)]
            uspub_refs = pd.concat([uspub_refs,ids])
            i += 1
        except KeyError:
            break
    #
    #
    #
    uspub_refs = uspub_refs[uspub_refs['Ref No'].map(lambda x: x==x)]
    uspub_refs = uspub_refs.reset_index(drop=True)
    uspub_refs.drop(uspub_refs.head(1).index,inplace=True)
    uspub_refs['Ref No'] = uspub_refs['Ref No'].map(lambda x: re.sub('\\.0','',str(x)))
    uspub_refs = uspub_refs.sort_values(by=['Ref No']).reset_index(drop=True)
    #uspub_refs['Date'] = uspub_refs['Date'].map(lambda x: datetime.strptime(str(x),'%Y-%m-%d').strftime('%B %Y'))
    formatted_refs = pd.concat([uspat_refs, uspub_refs])
    formatted_refs['Ref No'] = formatted_refs['Ref No'].map(lambda x: str(x))
    #formatted_refs['Name'] = formatted_refs['Name'].map(lambda x: x.title())
    #formatted_refs['Combined'] = formatted_refs['Ref No'].str.cat(formatted_refs[['Date','Name']], sep =' / ')
    formatted_refs = formatted_refs['Ref No'].to_string(index=False,header=False)
    bip_us_docs = formatted_refs.split('\n')
    # Foreign documents
    forn_docs = refs[forn_doc_cols]
    formatted_refs = pd.DataFrame([['no','date','co']],columns=['Ref No', 'Date', 'Country'])
    i = 0
    while True:
        # The spreadsheet has a column for each IDS/ref combo.  Iterate over the IDSs
        # sequentially until no more are left and then exit the loop
        doc_no =  'us-ids[0].us-foreign-document-cite[0].us-foreign-doc-reference[{}].doc-number[0]'.format(i)
        country =  'us-ids[0].us-foreign-document-cite[0].us-foreign-doc-reference[{}].country[0]'.format(i)
        date =  'us-ids[0].us-foreign-document-cite[0].us-foreign-doc-reference[{}].date[0]'.format(i)
        try:
            ids = forn_docs[[doc_no, date, country]]
            ids = ids.rename(columns={doc_no:'Ref No', date:'Date', country:'Country'})
            # remove blank rows
            ids = ids[ids.iloc[:,1].map(lambda x: x==x)]
            formatted_refs = pd.concat([formatted_refs,ids])
            i += 1
        except KeyError:
            break
    # Now some formatting to sort, drop blanks, reformat the date, and then make to string
    formatted_refs = formatted_refs.sort_values(by=['Country','Date']).reset_index(drop=True)
    formatted_refs.drop(formatted_refs.tail(1).index,inplace=True)
    formatted_refs['Date'] = formatted_refs['Date'].map(lambda x: datetime.strptime(str(x),'%Y-%m-%d').strftime('%b %Y'))
    formatted_refs['Ref No'] = formatted_refs['Ref No'].map(lambda x: re.sub('\\.0','',str(x)))
    formatted_refs['Ref No'] = formatted_refs['Ref No'].map(lambda x: re.sub('\\/','',str(x)))
    formatted_refs['Combined'] = formatted_refs['Country'].str.cat(formatted_refs['Ref No'])
    formatted_refs = formatted_refs['Combined'].to_string(index=False,header=False)
    bip_forn_docs = formatted_refs.split('\n')
    # NPL Docs
    npl_docs = refs[npl_doc_cols]
    formatted_refs = pd.DataFrame([['']],columns=['Doc Name'])
    i = 0
    while True:
        doc_name =  'us-ids[0].us-nplcit[{}].text[0]'.format(i)
        try:
            ids = npl_docs[[doc_name]]
            ids = ids.rename(columns={doc_name:'Doc Name'})
            formatted_refs = pd.concat([formatted_refs,ids])
            i += 1
        except KeyError:
            break
    formatted_refs = formatted_refs[formatted_refs['Doc Name'].map(lambda x: x==x)]
    formatted_refs = formatted_refs.reset_index(drop=True)
    formatted_refs.drop(formatted_refs.head(1).index,inplace=True)
    formatted_refs['Doc Name'] = formatted_refs['Doc Name'].map(lambda x: x.title())
    formatted_refs = formatted_refs.sort_values(by=['Doc Name']).reset_index(drop=True)
    formatted_refs = formatted_refs['Doc Name']
    formatted_refs = formatted_refs.tolist()
    formatted_refs = [x.replace('\r',' ').replace('\n',' ') for x in formatted_refs]
    formatted_refs = '\n'.join(formatted_refs)
    formatted_refs = formatted_refs.replace('January','Jan.').replace('February','Feb.').replace('March','Mar.')
    formatted_refs = formatted_refs.replace('April','Apr.').replace('May','May').replace('June','Jun.')
    formatted_refs = formatted_refs.replace('July','Jul.').replace('August','Aug.').replace('September','Sep.')
    formatted_refs = formatted_refs.replace('October','Oct.').replace('November','Nov.').replace('December','Dec.')
    formatted_refs = formatted_refs.replace('Patent Application','Appl.')
    bip_npl_docs = formatted_refs.split('\n')
    # new_docx('\n'.join(bip_us_docs) + '\n\n\n\n\n\n\n\n\n' + '\n'.join(bip_forn_docs) + '\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n' + '\n\n'.join(bip_npl_docs),pat_no + ' - BIP References.docx')
    #references = '\n'.join(bip_us_docs) + '\n' + '\n'.join(bip_forn_docs) + '\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n' + '\n\n'.join(bip_npl_docs)
    files = glob.glob('./*.docx')
    path = os.getcwd()
    for file in files:
        if 'Claim' in file and 'PTO' not in file:
            text = get_text(file)
            # new_docx(text,pat_no + ' - BIP Claims.docx')
            claims = text
        if 'Specification' in file and 'PTO' not in file:
            text = get_text(file)
            text = re.sub(u'[\u201c\u201d]','"',text)
            text = re.sub('’','\'',text)
            text = re.sub(' \t','',text)
            # new_docx(text,pat_no + ' - BIP Description.docx')
            description = text
            try:
                claims_part = re.search('Conclusion\n.+\n([\w\W]*)Abstract',description,re.IGNORECASE).group(1)
            except AttributeError:
                try:
                    claims_part = re.search('Claims\n([\w\W]*)Abstract',description,re.IGNORECASE).group(1)
                except AttributeError:
                    try:
                        claims_part = re.search('What is claimed is:\n([\w\W]*)Abstract',description,re.IGNORECASE).group(1)
                    except AttributeError:
                        claims_part = 'asdfasdfasdfa'
            description = description.replace(claims_part,'\n')
    return claims, description, bip_us_docs, bip_forn_docs, bip_npl_docs


################################################################################
################################################################################
                # Compare the docs
################################################################################

def compare_docs(file_1, file_2, matter_no):
    path = os.getcwd()
    file_1_dir = '.\\' + file_1
    file_2_dir = '.\\' + file_2
    # compare('.\\BIP Claims.docx', '.\\PTO Claims.docx','Compare Claims.docx')
    Application = win32com.client.gencache.EnsureDispatch("Word.Application")
    file_1_doc = Application.Documents.Open(os.getcwd() + file_1_dir)
    file_2_doc = Application.Documents.Open(os.getcwd() + file_2_dir)
    comp = Application.CompareDocuments(file_1_doc, file_2_doc)
    Application.ActiveDocument.ActiveWindow.View.Type = 3
    comp.SaveAs(FileName = path + "\\" +  date.today().strftime('%Y-%m-%d') + ' - ' +  matter_no + ' - Comparison.docx')
    comp.Close()
    Application.Quit()
    return

################################################################################
################################################################################
                # Generate Client's letter
################################################################################

def client_letter():
    if '1120' in os.getcwd():
        doc = Document('C:\\Users\\KellyThompson\\Dropbox (Brient IP)\\Attorneys\\Open Items\\7. Post Issuance\\Templates\\1120 OneTrust - Patent Report Certificate Template.docx')
    elif '1024' in os.getcwd():
        doc = Document('C:\\Users\\KellyThompson\\Dropbox (Brient IP)\\Attorneys\\Open Items\\7. Post Issuance\\Templates\\1024 Mohawk - Patent Report Certificate Template.docx')
    else:
        doc = Document('C:\\Users\\KellyThompson\\Dropbox (Brient IP)\\Attorneys\\Open Items\\7. Post Issuance\\Templates\\Other Clients - Patent Report Certificate Template.docx')
    current_path = os.getcwd()
    os.chdir('C:\\Users\\KellyThompson\\Dropbox (Brient IP)\\Attorneys\\Open Items\\7. Post Issuance\\Patent Cert Letters for Signature and Mailing')
    # Let's get the title
    r = str(soup).split('<meta content="patent" name="DC.type"/>')[1]
    s = r.split('" name="DC.title"/>')[0]
    title = s.split('<meta content="')[1]
    # Let's get issue date and maintenance dates
    x = str(soup).split('name="citation_patent_number"/>')[1]
    y = x.split('" name="DC.date" scheme="issue"/>')[0]
    issue_date = y.split('<meta content="')[1]    
    issue_date = datetime.strptime(str(issue_date),'%Y-%m-%d')
    date1 = issue_date + relativedelta(years = 3) + relativedelta(months=6)
    date2 = issue_date + relativedelta(years = 7) + relativedelta(months=6)
    date3 = issue_date + relativedelta(years = 11) + relativedelta(months=6)
    issue_date = issue_date.strftime('%B %d, %Y')
    date1 = date1.strftime('%B %d, %Y')
    date2 = date2.strftime('%B %d, %Y')
    date3 = date3.strftime('%B %d, %Y')
    matter_no = re.search('.+ - (.+) - .+',current_path).group(1)
    for para in doc.paragraphs:
        for run in para.runs:
            run.text
            if 'PATENTNO' in run.text:
                run.text = run.text.replace('PATENTNO',pat_no[0:2]+','+pat_no[2:5]+','+pat_no[5:])
            if 'TITLE' in run.text:
                run.text = run.text.replace('TITLE',title)
            if 'ISSUEDATE' in run.text:
                run.text = run.text.replace('ISSUEDATE',issue_date)
            if 'MATTERNO' in run.text:
                run.text = run.text.replace('MATTERNO',matter_no)
            if 'DATE1' in run.text:
                run.text = run.text.replace('DATE1',date1)
            if 'DATE2' in run.text:
                run.text = run.text.replace('DATE2',date2)
            if 'DATE3' in run.text:
                run.text = run.text.replace('DATE3',date3)
    doc.save(date.today().strftime('%Y-%m-%d') + ' - ' + matter_no + ' - Patent Report Certificate.docx')
    os.chdir(current_path)
    return

################################################################################
                # Function to copy over documents for analysis
################################################################################

def collect_docs(empty_folder):
    matter_no = re.search(' - (.+) - ',empty_folder).group(1)
    client_no = re.search('(.+?)-',matter_no).group(1)
    # ref_no =  re.search('-(.+)',matter_no).group(1)
    os.chdir('C:\\Users\\KellyThompson\\Dropbox (Brient IP)\\Files')
    for folder in next(os.walk('.'))[1]:
        if client_no in folder:
            os.chdir(os.getcwd() + '\\' + folder)
            break
    for folder in next(os.walk('.'))[1]:
        if matter_no in folder:
            os.chdir(os.getcwd() + '\\' + folder)
            break
    os.mkdir('C:\\Users\\KellyThompson\\Dropbox (Brient IP)\\Attorneys\\Open Items\\7. Post Issuance\\' + empty_folder + '\\IDS Forms')
    for subdir, dirs, files in os.walk(os.getcwd()):
        for file in files:
            filepath = subdir + os.sep + file
            # print(filepath)
            if 'Allowed' in filepath and 'docx' in filepath:
                shutil.copy2(filepath,'C:\\Users\\KellyThompson\\Dropbox (Brient IP)\\Attorneys\\Open Items\\7. Post Issuance\\' + empty_folder)
                # print('COPIED')
            elif 'IDS' in filepath and 'pdf' in filepath and 'Receipt' not in filepath and 'efiling' not in filepath:
                shutil.copy2(filepath,'C:\\Users\\KellyThompson\\Dropbox (Brient IP)\\Attorneys\\Open Items\\7. Post Issuance\\' + empty_folder + '\\IDS Forms')
                # print('COPIED')
            elif 'Specification' in filepath and 'docx' in filepath:
                shutil.copy2(filepath,'C:\\Users\\KellyThompson\\Dropbox (Brient IP)\\Attorneys\\Open Items\\7. Post Issuance\\' + empty_folder)
                # print('COPIED')



################################################################################
################################################################################
                # Now the main program to use these functions
                # idea is for this to be in the "Post Issuance" folder
################################################################################
os.chdir('C:\\Users\\KellyThompson\\Dropbox (Brient IP)\\Attorneys\\Open Items\\7. Post Issuance')
for folder in next(os.walk('.'))[1]:
    os.chdir('C:\\Users\\KellyThompson\\Dropbox (Brient IP)\\Attorneys\\Open Items\\7. Post Issuance')
    #print(folder)
    try:
        if 'Pat.' in folder:
            print(folder)
            global pat_no
            os.chdir(os.getcwd() + '\\' + folder)
            pat_no = os.getcwd()
            pat_no = re.search('.+Pat\. (.+)',pat_no).group(1).replace(',','')
            #flag1 = any('Claim' in file for file in os.listdir()) # True if there's a file called claims - i.e., ready for analysis
            #flag2 = any(('Specification' in file and 'pdf' not in file) for file in os.listdir()) # True if there's a file called specification - i.e., ready for analysis
            #flag3 = any('IDS Forms' in file for file in os.listdir()) # True if there's a folder called IDS Forms - i.e., ready for analysis
            flag4 = any('Comparison' in file for file in os.listdir()) # True if there's a file called comparison - i.e., anaysis is complete
            #if flag1 == False or flag2 == False or flag3 == False: # so the folder does not have a claim AND a spec file
             #   try:
                    #collect_docs(folder)
              #      print('File isn\'t ready, please copy files over. Confirm and then re-run for: ' + pat_no)
               #     os.chdir('..')
                #    continue
                #except:
                 #   print('File isn\'t ready, missing either spec, claims, or IDS folder: Pat. No. ' + pat_no)
                 #   os.chdir('..')
                  #  continue
            if flag4:
                print('Found already completed Pat. No. ' + pat_no)
                os.chdir('..')
                continue
            matter_no = re.search('.+ - (.+) - .+',os.getcwd()).group(1)
            [p_claims, p_description, p_us_ref, p_for_ref, p_url_ref] = pto_docs(pat_no)
            [b_claims, b_description, b_us_ref, b_for_ref, p_url_ref] = bip_docs()
            pto_doc = date.today().strftime('%Y-%m-%d') + ' - ' + matter_no + ' - PTO Printed Docs.docx'
            bip_doc = date.today().strftime('%Y-%m-%d') + ' - ' + matter_no + ' - BIP As-Filed Docs.docx'
            ref_doc = date.today().strftime('%Y-%m-%d') + ' - ' + matter_no + ' - References Comparison.xlsx'
            new_docx(pto_doc, p_description, p_claims)
            new_docx(bip_doc, b_description, b_claims)
            new_spreadsheet(ref_doc, p_us_ref, p_for_ref, p_url_ref, b_us_ref, b_for_ref, p_url_ref)
            compare_docs(bip_doc, pto_doc, matter_no)
            client_letter()
            os.chdir('..')
            print('Finished Pat. No. ' + pat_no)
    except Exception as e:
        print('Problem with Pat. No. ' + pat_no)
        os.chdir('..')
        print(e)
